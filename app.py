import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
import os
import io
from docx import Document

# LangChain Neo4j kütüphaneleri
from langchain_community.graphs import Neo4jGraph
from langchain_community.chains.graph_qa.cypher import GraphCypherQAChain

# --- SAYFA YAPILANDIRMASI ---
st.set_page_config(page_title="EOS | Epistemic OS (Edge Weighting)", page_icon="🧠", layout="wide")
st.title("🏛️ Epistemic Operating System (EOS) v3.0 - Dinamik Kenar Ağırlığı Aktif")
st.markdown("*Multi-Layered Ontology & Dynamic Edge Weighting Engine*")

# --- YAN MENÜ (İŞLETİM SİSTEMİ KONTROLLERİ) ---
with st.sidebar:
    st.header("⚙️ EOS Kontrol Paneli")
    
    # Güvenli kasadan anahtarları çekme
    try:
        os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]
    except Exception:
        pass
    
    # Neo4j Bağlantı Kontrolü
    neo4j_connected = False
    try:
        NEO4J_URI = st.secrets["NEO4J_URI"]
        NEO4J_USER = st.secrets["NEO4J_USER"]
        NEO4J_PASSWORD = st.secrets["NEO4J_PASSWORD"]
        
        graph = Neo4jGraph(url=NEO4J_URI, username=NEO4J_USER, password=NEO4J_PASSWORD)
        neo4j_connected = True
        st.sidebar.success("Knowledge Graph & Ontoloji Bağlı 🕸️")
    except Exception as e:
        st.sidebar.warning(f"Graf bağlanamadı: {str(e)}")
    
    # 3. AŞAMA: DİNAMİK KENAR AĞIRLIKLANDIRMA (YENİ EKLENDİ)
    st.markdown("---")
    st.subheader("🔗 İlişki Ağırlığı (Edge Weight)")
    edge_weight = st.slider(
        "Minimum Bağlantı Şiddeti", 
        min_value=1, max_value=10, value=5,
        help="1: Tüm dolaylı ve zayıf ihtimaller dahil. 10: Sadece doğrudan, kesin ve çok güçlü kanıtlanmış etkileşimler."
    )
    
    st.markdown("---")
    
    # Rol Seçimi
    active_role = st.selectbox(
        "🎭 Aktif Yapay Zekâ Rolü",
        ["Moderatör", "Sokratik Öğretmen", "Şeytanın Avukatı", "Kör Nokta Avcısı", "Şarih", "Hakem", "Köprü Kurucu", "Araştırma Danışmanı"]
    )
    
    # Mod Seçimi
    active_mode = st.radio(
        "🔄 Epistemik Mod",
        ["Standart Analiz", "Tarihsel Mod (Zaman Kilidi)", "Modern Mod (Karşılaştırmalı)", "Kavram Evrimi", "Tartışma Sonu Raporu"]
    )
    
    year_lock = None
    if active_mode == "Tarihsel Mod (Zaman Kilidi)":
        year_lock = st.number_input("Tarih Kilidi (Örn: 1198)", value=1198)

# --- SİSTEM PROMPTU VE DİNAMİK KENAR AĞIRLIK ÇERÇEVESİ ---
def build_system_prompt(role, mode, year, weight):
    base_prompt = f"""
    # MİSYON VE ONTOLOJİK ÇERÇEVE
    Sen sıradan bir yapay zekâ sohbet botu değil, eser merkezli çalışan çok katmanlı bir 'Epistemik İşletim Sistemi'sin.
    Bilgiyi işlerken şu katmanlı ontolojik hiyerarşiyi baz alırsın:
    1. Ontolojik Düzey (Metafizik, Epistemoloji, Siyaset, Ahlak vb.)
    2. Filozof ve Eser Matrisi
    3. Kavramsal Evrim
    
    # DİNAMİK KENAR AĞIRLIKLANDIRMA KURALI (ÖNEMLİ):
    Kullanıcı şu an sistemin 'Minimum Bağlantı Şiddetini' 10 üzerinden {weight} olarak belirledi. 
    - Eğer bu değer (7-10) arasındaysa: Yalnızca doğrudan, birincil kaynaklarla kanıtlanmış, çok güçlü nedensellik barındıran ilişkileri kabul et. Zayıf ihtimalleri reddet.
    - Eğer bu değer (4-6) arasındaysa: Standart akademik etkileşimleri ve dolaylı atıfları kabul et.
    - Eğer bu değer (1-3) arasındaysa: Spekülatif, zayıf, gizli kalmış ve yan etkileşimleri (kelebek etkisi) gün yüzüne çıkar ve varsayımsal köprüler kur.
    
    Amacın kullanıcının yerine düşünmek değil; düşünmesini sağlamak ve argümanların ontolojik zeminini bu ağırlık kuralına göre sorgulatmaktır.
    """
    
    role_instructions = {
        "Sokratik Öğretmen": "Asla doğrudan cevap verme. Yalnızca 'Neden?', 'Bu iddianın ontolojik dayanağı nedir?' gibi sorular üret.",
        "Şeytanın Avukatı": "Kullanıcının savunduğu görüşün tam tersini ontolojik argümanlarla en güçlü haliyle savun.",
        "Kör Nokta Avcısı": "Kullanıcının fark etmediği kavramsal kaymaları ve eksik bağlamları tespit et.",
        "Şarih": "Karmaşık kavramları katmanlarına ayırarak açıkla ama orijinal terimleri koru.",
        "Köprü Kurucu": "Farklı dönemler ve ontolojik katmanlar arasında köprü kur.",
        "Araştırma Danışmanı": "Oturum sonunda literatürde az çalışılmış ontolojik problemler üret.",
        "Moderatör": "Tartışmayı yönet, tarafsız kal ve konuyu toparla.",
        "Hakem": "Hangi iddiaların ontolojik olarak temellendirildiğini ve hangi soruların cevapsız kaldığını göster."
    }
    return f"{base_prompt}\n\n# ŞU ANKİ ROLÜN: {role}\n{role_instructions.get(role, '')}"

# --- SOHBET GEÇMİŞİ YÖNETİMİ ---
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message("user" if isinstance(msg, HumanMessage) else "assistant"):
        st.markdown(msg.content)

# --- KULLANICI GİRİŞİ VE GRAF SORGULAMA ---
if prompt := st.chat_input("Düşünceyi başlat (Örn: 'Gazali'nin İbn Sina eleştirisinin ana damarı...'"):
    st.session_state.messages.append(HumanMessage(content=prompt))
    with st.chat_message("user"):
        st.markdown(prompt)
        
    with st.chat_message("assistant"):
        try:
            llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0.7)
            
            # --- ÇOK KATMANLI GRAFTAN BİLGİ ÇEKME ---
            graph_context = ""
            if neo4j_connected:
                with st.spinner("Kenar Ağırlıkları Hesaplanıyor ve Graf Taranıyor... 🕸️⚖️"):
                    try:
                        cypher_chain = GraphCypherQAChain.from_llm(
                            cypher_llm=llm,
                            qa_llm=llm,
                            graph=graph,
                            verbose=False,
                            return_direct=True
                        )
                        graph_result = cypher_chain.invoke({"query": prompt})["result"]
                        if graph_result:
                            graph_context = f"\n\n# KNOWLEDGE GRAPH BAĞLANTI VERİLERİ:\n{graph_result}"
                    except Exception:
                        pass
            
            # --- YANIT ÜRETME ---
            final_system_prompt = build_system_prompt(active_role, active_mode, year_lock, edge_weight) + graph_context
            sys_msg = SystemMessage(content=final_system_prompt)
            api_messages = [sys_msg] + st.session_state.messages
            
            with st.spinner("Epistemik motor bağlantı şiddetini analiz ediyor... 🧠"):
                response = llm.invoke(api_messages)
                
            st.markdown(response.content)
            st.session_state.messages.append(AIMessage(content=response.content))
            
        except Exception as e:
            st.error(f"Bir hata oluştu: {str(e)}")

# --- WORD DOSYASI İNDİRME ÖZELLİĞİ ---
def generate_word_document(messages):
    doc = Document()
    doc.add_heading('Epistemik İşletim Sistemi - Çok Katmanlı Ontoloji Raporu', 0)
    
    for msg in messages:
        if isinstance(msg, HumanMessage):
            doc.add_heading('Araştırmacı:', level=2)
            doc.add_paragraph(msg.content)
        elif isinstance(msg, AIMessage):
            doc.add_heading('Epistemik Motor:', level=2)
            doc.add_paragraph(msg.content)
        doc.add_paragraph('_' * 40)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

if st.session_state.messages:
    st.sidebar.markdown("---")
    st.sidebar.subheader("📄 Dışa Aktar")
    word_data = generate_word_document(st.session_state.messages)
    st.sidebar.download_button(
        label="Tüm Diyaloğu Word Olarak İndir",
        data=word_data,
        file_name="epistemik_ontoloji_diyalog.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )